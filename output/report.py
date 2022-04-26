import os

from PyQt5.QtWidgets import QWidget, QFileDialog, QMessageBox, QTableWidgetItem
from docx.shared import Cm, Pt

from .ui_report import Ui_DialogReport

from openpyxl import Workbook
from openpyxl.styles import Alignment, Side, Border, Font
import docx


class ReportWindow(QWidget):
    def __init__(self, data_table):
        super(ReportWindow, self).__init__()
        self.ui = Ui_DialogReport()
        self.ui.setupUi(self)

        self.dir = ''
        self.data = data_table

        self.update_tabel()

        self.ui.btn_chouse_dir.clicked.connect(self.get_dir)
        self.ui.btn_report.clicked.connect(self.report)
        self.ui.btn_cancel.clicked.connect(self.close)

        self.show()

    def get_dir(self):
        selected_dir = QFileDialog.getExistingDirectory(self, caption='Choose Directory', directory=os.getcwd())
        if selected_dir:
            self.dir = selected_dir
            self.ui.le_folder_path.setText(selected_dir)

    def update_tabel(self):
        for i in range(len(self.data)):
            self.ui.tableWidget.setItem(i + 1, 2, QTableWidgetItem(f"{self.data[i][2]}"))

    def report(self):
        if not self.dir:
            self.showDialog("Ошибка", "Ошибка: Выберите путь.")
        elif not (self.ui.checkBox_word.isChecked() or self.ui.checkBox_excel.isChecked()):
            self.showDialog("Ошибка", "Ошибка: Выберите нужный формат файла.")
        else:
            if self.ui.checkBox_excel.isChecked():
                self.create_excel(self.data)
            if self.ui.checkBox_word.isChecked():
                self.create_word(self.data)
            if self.ui.checkBox_word.isChecked() and self.ui.checkBox_excel.isChecked():
                self.create_excel(self.data)
                self.create_word(self.data)
            self.showDialog("Успех", "Отчет создан.")
            self.close()

    def create_excel(self, items):
        # Создаем новый лист
        workbook = Workbook()
        ws = workbook.active

        # Стили шрифта
        font_title = Font(name='Times New Roman', size=14, bold=True)
        font_table = Font(name='Times New Roman', size=12)

        # Заголовок таблицы
        ws['A1'] = '№\nп/п'
        ws['B1'] = 'Наименование показателя'
        ws['C1'] = 'Значение показателя'

        # Горизонтальное и вертикальное выравнивание текста
        ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
        ws['B1'].alignment = Alignment(horizontal='center', vertical='center')
        ws['C1'].alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

        # Оформление границ ячеек
        ws.row_dimensions[1].height = 33
        ws.column_dimensions['B'].width = 60
        ws.column_dimensions['C'].width = 15

        ws.row_dimensions[3].height = 30
        ws.row_dimensions[9].height = 30
        ws.row_dimensions[10].height = 30

        for i in range(4, 9):
            ws.row_dimensions[i].height = 40

        # Стиль текста ячеек
        ws['A1'].font = font_title
        ws['B1'].font = font_title
        ws['C1'].font = font_title

        # Заполняем таблицу
        row = 2

        for inum, text, value in items:
            ws[f'A{row}'] = inum
            ws[f'A{row}'].font = font_table
            ws[f'A{row}'].alignment = Alignment(horizontal='center', vertical='center')

            ws[f'B{row}'] = text
            ws[f'B{row}'].font = font_table
            ws[f'B{row}'].alignment = Alignment(
                horizontal='general',
                vertical='bottom',
                text_rotation=0,
                wrap_text=True,
                shrink_to_fit=False,
                indent=0
            )

            ws[f'C{row}'] = value
            ws[f'C{row}'].font = font_table
            ws[f'C{row}'].alignment = Alignment(horizontal='right')

            row += 1

        # Устанавливаем границы таблицы
        thin = Side(border_style="thin", color="000000")
        for row in ws[f'A1:C{len(items) + 1}']:
            for cell in row:
                cell.border = Border(top=thin, left=thin, right=thin, bottom=thin)

        # Сохраняем файл
        workbook.save(f"{self.dir}/Образовательная деятельность СПО Excel.xlsx")

    def create_word(self, items):
        # Create an instance of a word document
        document = docx.Document()

        # Add a Title to the document
        document.add_heading('Шапка', 0)

        # добавляем таблицу с одной строкой
        # для заполнения названий колонок
        table = document.add_table(1, len(items[0]))
        # определяем стиль таблицы
        table.allow_autofit = True
        table.style = 'Table Grid'
        # Получаем строку с колонками из добавленной таблицы
        head_cells = table.rows[0].cells
        # добавляем названия колонок
        for i, item in enumerate(['  №\nП/П', 'Наименование показателя', 'Значение показателя']):
            p = head_cells[i].paragraphs[0]
            # название колонки
            p.add_run(item).bold = True

        # добавляем данные к существующей таблице
        for row in items:
            # добавляем строку с ячейками к объекту таблицы
            cells = table.add_row().cells
            for i, item in enumerate(row):
                # вставляем данные в ячейки
                cells[i].text = str(item)
                # Шрифт и размер
                cells[i].paragraphs[0].runs[0].font.name = 'Times New Roman'
                cells[i].paragraphs[0].runs[0].font.size = Pt(12)

        table.columns[0].width = Cm(1.5)
        table.columns[1].width = Cm(11.5)
        table.columns[2].width = Cm(3.3)

        # Now save the document to a location
        document.save(f'{self.dir}/Образовательная деятельность СПО Word.docx')

    def showDialog(self, title, text):
        msgBox = QMessageBox()
        msgBox.setIcon(QMessageBox.Information)
        msgBox.setText(text)
        msgBox.setWindowTitle(title)
        msgBox.setStandardButtons(QMessageBox.Ok)
        msgBox.exec()


